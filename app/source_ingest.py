from __future__ import annotations

import json
import tempfile
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Optional, Sequence

import requests

from app.model_api import chat_completion_text
from app.settings import Settings


SUPPORTED_SOURCE_SUFFIXES = {".txt", ".md", ".pdf", ".docx"}
MAX_SOURCE_FILES = 5
SOURCE_CHUNK_SIZE = 5000
MAX_CHUNKS_PER_SOURCE = 24
MAX_REFINED_CHUNK_CHARS = 2400


@dataclass(frozen=True)
class SourceFileInput:
    name: str
    data: bytes


@dataclass(frozen=True)
class SourceIngestRuntimeConfig:
    text_provider: str
    text_base_url: str
    text_api_key: str
    text_model: str
    mineru_base_url: str
    mineru_api_key: str
    mineru_model_version: str
    mineru_language: str
    mineru_enable_formula: bool
    mineru_enable_table: bool
    mineru_is_ocr: bool
    mineru_poll_interval_seconds: float
    mineru_timeout_seconds: int


@dataclass(frozen=True)
class ExtractedSource:
    name: str
    suffix: str
    text: str
    extraction_method: str
    metadata: dict[str, Any]


@dataclass(frozen=True)
class PreparedRequirement:
    final_requirement: str
    extracted_sources: list[ExtractedSource]
    synthesis_summary: str


class SourceDocumentProcessor:
    def __init__(self, settings: Settings) -> None:
        self.settings = settings

    def build_runtime_config(
        self,
        *,
        text_provider: Optional[str] = None,
        text_base_url: Optional[str] = None,
        text_api_key: Optional[str] = None,
        text_model: Optional[str] = None,
        mineru_base_url: Optional[str] = None,
        mineru_api_key: Optional[str] = None,
        mineru_model_version: Optional[str] = None,
        mineru_language: Optional[str] = None,
        mineru_enable_formula: Optional[bool] = None,
        mineru_enable_table: Optional[bool] = None,
        mineru_is_ocr: Optional[bool] = None,
        mineru_poll_interval_seconds: Optional[float] = None,
        mineru_timeout_seconds: Optional[int] = None,
    ) -> SourceIngestRuntimeConfig:
        cfg = SourceIngestRuntimeConfig(
            text_provider=(text_provider or self.settings.text_provider or "openai").strip().lower(),
            text_base_url=(text_base_url or self.settings.text_base_url).strip(),
            text_api_key=(text_api_key or self.settings.text_api_key).strip(),
            text_model=(text_model or self.settings.text_model).strip(),
            mineru_base_url=(mineru_base_url or self.settings.resolved_mineru_base_url).strip(),
            mineru_api_key=(mineru_api_key or self.settings.resolved_mineru_api_key).strip(),
            mineru_model_version=(mineru_model_version or self.settings.mineru_model_version).strip(),
            mineru_language=(mineru_language or self.settings.mineru_language).strip(),
            mineru_enable_formula=bool(
                self.settings.mineru_enable_formula if mineru_enable_formula is None else mineru_enable_formula
            ),
            mineru_enable_table=bool(
                self.settings.mineru_enable_table if mineru_enable_table is None else mineru_enable_table
            ),
            mineru_is_ocr=bool(self.settings.mineru_is_ocr if mineru_is_ocr is None else mineru_is_ocr),
            mineru_poll_interval_seconds=float(
                self.settings.mineru_poll_interval_seconds
                if mineru_poll_interval_seconds is None
                else mineru_poll_interval_seconds
            ),
            mineru_timeout_seconds=int(
                self.settings.mineru_timeout_seconds if mineru_timeout_seconds is None else mineru_timeout_seconds
            ),
        )
        if cfg.text_provider not in {"openai", "gemini"}:
            raise ValueError("Source ingestion text provider must be `openai` or `gemini`.")
        if not cfg.text_base_url:
            raise ValueError("Source ingestion text base_url cannot be empty.")
        if not cfg.text_api_key:
            raise ValueError("Source ingestion text api_key cannot be empty.")
        if not cfg.text_model:
            raise ValueError("Source ingestion text model cannot be empty.")
        if cfg.mineru_poll_interval_seconds <= 0:
            raise ValueError("mineru_poll_interval_seconds must be > 0.")
        if cfg.mineru_timeout_seconds < 30:
            raise ValueError("mineru_timeout_seconds must be >= 30.")
        return cfg

    def prepare_requirement(
        self,
        *,
        user_requirement: str,
        source_files: Sequence[SourceFileInput],
        runtime_cfg: SourceIngestRuntimeConfig,
    ) -> PreparedRequirement:
        cleaned_requirement = (user_requirement or "").strip()
        if not source_files:
            return PreparedRequirement(
                final_requirement=cleaned_requirement,
                extracted_sources=[],
                synthesis_summary="",
            )
        if len(source_files) > MAX_SOURCE_FILES:
            raise ValueError(f"At most {MAX_SOURCE_FILES} source files are supported per request.")

        extracted_sources: list[ExtractedSource] = []
        with tempfile.TemporaryDirectory(prefix="editdeck_sources_") as temp_dir_raw:
            temp_dir = Path(temp_dir_raw)
            for index, source_file in enumerate(source_files, start=1):
                extracted_sources.append(
                    self._extract_source_file(
                        source_file=source_file,
                        runtime_cfg=runtime_cfg,
                        temp_dir=temp_dir / f"source_{index:02d}",
                    )
                )

        final_requirement, summary = self._synthesize_requirement(
            user_requirement=cleaned_requirement,
            extracted_sources=extracted_sources,
            runtime_cfg=runtime_cfg,
        )
        return PreparedRequirement(
            final_requirement=final_requirement,
            extracted_sources=extracted_sources,
            synthesis_summary=summary,
        )

    def _extract_source_file(
        self,
        *,
        source_file: SourceFileInput,
        runtime_cfg: SourceIngestRuntimeConfig,
        temp_dir: Path,
    ) -> ExtractedSource:
        temp_dir.mkdir(parents=True, exist_ok=True)
        safe_name = Path(source_file.name or "source").name
        suffix = Path(safe_name).suffix.lower()
        if suffix not in SUPPORTED_SOURCE_SUFFIXES:
            supported = ", ".join(sorted(SUPPORTED_SOURCE_SUFFIXES))
            raise ValueError(f"Unsupported source file `{safe_name}`. Supported types: {supported}.")

        file_path = temp_dir / safe_name
        file_path.write_bytes(source_file.data)

        if suffix in {".txt", ".md"}:
            text = self._read_text_file(file_path)
            return ExtractedSource(
                name=safe_name,
                suffix=suffix,
                text=text,
                extraction_method="plain_text",
                metadata={"char_count": len(text)},
            )

        if suffix == ".docx":
            text = self._extract_docx_text(file_path)
            return ExtractedSource(
                name=safe_name,
                suffix=suffix,
                text=text,
                extraction_method="python-docx",
                metadata={"char_count": len(text)},
            )

        text, metadata = self._extract_pdf_text_with_mineru(file_path=file_path, runtime_cfg=runtime_cfg, work_dir=temp_dir)
        return ExtractedSource(
            name=safe_name,
            suffix=suffix,
            text=text,
            extraction_method="mineru",
            metadata=metadata,
        )

    @staticmethod
    def _read_text_file(file_path: Path) -> str:
        encodings = ("utf-8", "utf-8-sig", "utf-16", "gb18030", "latin-1")
        for encoding in encodings:
            try:
                text = file_path.read_text(encoding=encoding)
                if encoding == "latin-1":
                    text = text.encode("latin-1").decode("utf-8", errors="ignore") or text
                return text.strip()
            except Exception:
                continue
        raise ValueError(f"Unable to decode text file: {file_path.name}")

    @staticmethod
    def _extract_docx_text(file_path: Path) -> str:
        try:
            from docx import Document
            from docx.document import Document as DocumentObject
            from docx.oxml.table import CT_Tbl
            from docx.oxml.text.paragraph import CT_P
            from docx.table import _Cell, Table
            from docx.text.paragraph import Paragraph
        except ImportError as exc:
            raise RuntimeError("DOCX extraction requires `python-docx`. Please install project dependencies.") from exc

        def iter_block_items(parent: Any):
            if isinstance(parent, DocumentObject):
                parent_elm = parent.element.body
            elif isinstance(parent, _Cell):
                parent_elm = parent._tc
            else:
                raise TypeError(f"Unsupported parent type: {type(parent)}")

            for child in parent_elm.iterchildren():
                if isinstance(child, CT_P):
                    yield Paragraph(child, parent)
                elif isinstance(child, CT_Tbl):
                    yield Table(child, parent)

        document = Document(str(file_path))
        blocks: list[str] = []

        def append_line(line: str) -> None:
            cleaned = " ".join((line or "").split()).strip()
            if cleaned:
                blocks.append(cleaned)

        for block in iter_block_items(document):
            if isinstance(block, Paragraph):
                style_name = getattr(getattr(block, "style", None), "name", "") or ""
                text = block.text.strip()
                if not text:
                    continue
                if style_name.lower().startswith("heading"):
                    append_line(f"[{text}]")
                else:
                    append_line(text)
                continue

            rows: list[str] = []
            for row in block.rows:
                cells = [" ".join(cell.text.split()).strip() for cell in row.cells]
                cells = [cell for cell in cells if cell]
                if cells:
                    rows.append(" | ".join(cells))
            if rows:
                append_line("[TABLE]")
                blocks.extend(rows)

        for section in document.sections:
            for paragraph in section.header.paragraphs:
                append_line(f"[HEADER] {paragraph.text}")
            for paragraph in section.footer.paragraphs:
                append_line(f"[FOOTER] {paragraph.text}")

        text = "\n".join(blocks).strip()
        if not text:
            raise ValueError(f"No readable text found in DOCX file: {file_path.name}")
        return text

    def _extract_pdf_text_with_mineru(
        self,
        *,
        file_path: Path,
        runtime_cfg: SourceIngestRuntimeConfig,
        work_dir: Path,
    ) -> tuple[str, dict[str, Any]]:
        if not runtime_cfg.mineru_api_key:
            raise ValueError("PDF extraction requires MINERU_API_KEY or --mineru-api-key.")

        session = requests.Session()
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {runtime_cfg.mineru_api_key}",
        }
        request_payload = {
            "enable_formula": runtime_cfg.mineru_enable_formula,
            "language": runtime_cfg.mineru_language,
            "enable_table": runtime_cfg.mineru_enable_table,
            "files": [
                {
                    "name": file_path.name,
                    "data_id": f"source-{file_path.stem}",
                    "is_ocr": runtime_cfg.mineru_is_ocr,
                }
            ],
            "model_version": runtime_cfg.mineru_model_version,
        }
        response = session.post(
            f"{runtime_cfg.mineru_base_url.rstrip('/')}/file-urls/batch",
            headers=headers,
            json=request_payload,
            timeout=min(runtime_cfg.mineru_timeout_seconds, 120),
        )
        payload = self._parse_mineru_response(response, "upload-url request")
        data = payload.get("data") or {}
        batch_id = str(data.get("batch_id") or "").strip()
        file_urls = data.get("file_urls") or []
        upload_url = str(file_urls[0] or "").strip() if isinstance(file_urls, list) and file_urls else ""
        if not batch_id or not upload_url:
            raise RuntimeError(f"MinerU upload-url response is missing batch_id or file_urls: {payload}")

        with file_path.open("rb") as handle:
            upload_response = session.put(upload_url, data=handle, timeout=min(runtime_cfg.mineru_timeout_seconds, 300))
        if upload_response.status_code not in {200, 201}:
            raise RuntimeError(
                f"MinerU upload failed with HTTP {upload_response.status_code}: {upload_response.text[:300]}"
            )

        result_payload = self._poll_mineru_batch_result(
            session=session,
            headers=headers,
            base_url=runtime_cfg.mineru_base_url,
            batch_id=batch_id,
            timeout_seconds=runtime_cfg.mineru_timeout_seconds,
            poll_interval_seconds=runtime_cfg.mineru_poll_interval_seconds,
        )
        extract_rows = ((result_payload.get("data") or {}).get("extract_result") or [])
        if not isinstance(extract_rows, list) or not extract_rows:
            raise RuntimeError(f"MinerU batch result does not contain extract_result: {result_payload}")
        result_row = extract_rows[0]
        full_zip_url = str(result_row.get("full_zip_url") or "").strip()
        if not full_zip_url:
            raise RuntimeError(f"MinerU result does not contain full_zip_url: {result_payload}")

        zip_path = work_dir / "mineru_result.zip"
        extracted_dir = work_dir / "mineru_extracted"
        extracted_dir.mkdir(parents=True, exist_ok=True)
        download_response = session.get(full_zip_url, timeout=min(runtime_cfg.mineru_timeout_seconds, 300))
        download_response.raise_for_status()
        zip_path.write_bytes(download_response.content)
        with zipfile.ZipFile(zip_path, "r") as archive:
            archive.extractall(extracted_dir)

        text = self._extract_text_from_mineru_output(extracted_dir)
        metadata = {
            "batch_id": batch_id,
            "char_count": len(text),
            "model_version": runtime_cfg.mineru_model_version,
        }
        return text, metadata

    @staticmethod
    def _parse_mineru_response(response: requests.Response, action: str) -> dict[str, Any]:
        try:
            payload = response.json()
        except ValueError as exc:
            response.raise_for_status()
            raise RuntimeError(f"MinerU {action} did not return JSON.") from exc
        if response.status_code != 200:
            raise RuntimeError(f"MinerU {action} failed with HTTP {response.status_code}: {payload}")
        code = payload.get("code", -1)
        try:
            code_value = int(code)
        except (TypeError, ValueError):
            code_value = -1
        if code_value != 0:
            raise RuntimeError(f"MinerU {action} failed: {payload.get('msg') or payload}")
        return payload

    def _poll_mineru_batch_result(
        self,
        *,
        session: requests.Session,
        headers: dict[str, str],
        base_url: str,
        batch_id: str,
        timeout_seconds: int,
        poll_interval_seconds: float,
    ) -> dict[str, Any]:
        import time

        deadline = time.monotonic() + timeout_seconds
        last_payload: Optional[dict[str, Any]] = None
        while time.monotonic() < deadline:
            response = session.get(
                f"{base_url.rstrip('/')}/extract-results/batch/{batch_id}",
                headers=headers,
                timeout=min(timeout_seconds, 120),
            )
            payload = self._parse_mineru_response(response, "result polling")
            last_payload = payload
            rows = ((payload.get("data") or {}).get("extract_result") or [])
            if isinstance(rows, list) and rows:
                row = rows[0]
                state = str(row.get("state") or "").strip().lower()
                if state == "done":
                    return payload
                if state in {"failed", "error"}:
                    raise RuntimeError(f"MinerU parse failed: {row.get('err_msg') or row}")
            time.sleep(poll_interval_seconds)
        raise TimeoutError(f"MinerU parse timed out after {timeout_seconds} seconds. Last payload: {last_payload}")

    @staticmethod
    def _extract_text_from_mineru_output(extracted_dir: Path) -> str:
        markdown_candidates = sorted(extracted_dir.rglob("*.md"))
        if markdown_candidates:
            best_markdown = max(markdown_candidates, key=lambda path: path.stat().st_size)
            text = best_markdown.read_text(encoding="utf-8", errors="ignore").strip()
            if text:
                return text

        json_candidates = sorted(extracted_dir.rglob("*_content_list.json"))
        for candidate in json_candidates:
            try:
                rows = json.loads(candidate.read_text(encoding="utf-8"))
            except Exception:
                continue
            if not isinstance(rows, list):
                continue
            lines: list[str] = []
            for row in rows:
                if not isinstance(row, dict):
                    continue
                text = str(row.get("text") or row.get("content") or "").strip()
                if text:
                    lines.append(text)
            merged = "\n".join(lines).strip()
            if merged:
                return merged

        raise ValueError("MinerU output did not contain readable markdown or text JSON.")

    def _synthesize_requirement(
        self,
        *,
        user_requirement: str,
        extracted_sources: Sequence[ExtractedSource],
        runtime_cfg: SourceIngestRuntimeConfig,
    ) -> tuple[str, str]:
        source_blocks: list[str] = []
        for source in extracted_sources:
            raw_text = (source.text or "").strip()
            if not raw_text:
                continue
            refined_chunks = self._refine_source_for_ppt(
                user_requirement=user_requirement,
                source=source,
                runtime_cfg=runtime_cfg,
            )
            if not refined_chunks:
                continue
            source_blocks.append(
                "\n".join(
                    [
                        f"[Source File] {source.name}",
                        f"[Type] {source.suffix}",
                        f"[Extraction] {source.extraction_method}",
                        "[Relevant Extracts For PPT]",
                        "\n".join(refined_chunks),
                    ]
                )
            )

        if not source_blocks:
            return user_requirement, ""

        user_prompt = f"""
你是 PPT 策划前处理助手。你的任务是把“用户的一句话需求”和“附带文件里提取出的内容”融合，整理成一份可直接交给 PPT 生成流水线的需求说明。

输出必须是严格 JSON，且只能输出 JSON：
{{
  "final_requirement": "string",
  "summary": "string"
}}

要求：
1. final_requirement 使用中文，面向后续 PPT 生成模型，必须保留用户明确提出的目标、受众、语气、页数偏好或重点要求。
2. 你要主动吸收文件中的主题、章节结构、关键事实、数据结论、方案步骤、案例信息和专业术语，但不要原样堆砌全文。
3. final_requirement 应该像一份完整的 PPT 任务说明，既有主题，也有建议覆盖的核心内容范围。
4. 如果文件内容与用户一句话有冲突，以用户明确要求为最高优先级。
5. 不要编造文件里没有的信息；不确定时用更保守的概括。
6. summary 用 1-3 句概括你从文件里提炼了什么。
7. 不要输出 markdown。

[User Requirement]
{user_requirement or "用户未提供额外文字要求。"}

[Extracted File Content]
{chr(10).join(source_blocks)}
""".strip()
        system_prompt = "你负责把原始文件内容整理成稳定、清晰、适合生成 PPT 的任务说明。"
        raw = chat_completion_text(
            provider=runtime_cfg.text_provider,
            base_url=runtime_cfg.text_base_url,
            api_key=runtime_cfg.text_api_key,
            model=runtime_cfg.text_model,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            temperature=0.2,
        )
        parsed = self._parse_json_object(raw)
        final_requirement = str(parsed.get("final_requirement") or "").strip()
        summary = str(parsed.get("summary") or "").strip()
        if not final_requirement:
            raise ValueError("File synthesis step did not return final_requirement.")
        return final_requirement, summary

    def _refine_source_for_ppt(
        self,
        *,
        user_requirement: str,
        source: ExtractedSource,
        runtime_cfg: SourceIngestRuntimeConfig,
    ) -> list[str]:
        chunks = self._chunk_text(source.text, SOURCE_CHUNK_SIZE)
        if len(chunks) > MAX_CHUNKS_PER_SOURCE:
            chunks = chunks[:MAX_CHUNKS_PER_SOURCE]

        refined_chunks: list[str] = []
        for index, chunk in enumerate(chunks, start=1):
            refined = self._refine_chunk_for_ppt(
                user_requirement=user_requirement,
                source=source,
                chunk=chunk,
                chunk_index=index,
                total_chunks=len(chunks),
                runtime_cfg=runtime_cfg,
            )
            if refined:
                refined_chunks.append(refined)
        return refined_chunks

    def _refine_chunk_for_ppt(
        self,
        *,
        user_requirement: str,
        source: ExtractedSource,
        chunk: str,
        chunk_index: int,
        total_chunks: int,
        runtime_cfg: SourceIngestRuntimeConfig,
    ) -> str:
        user_prompt = f"""
你是 PPT 前处理分析助手。请判断下面这段文档内容，对“生成这份 PPT”是否真的有帮助。

输出必须是严格 JSON，且只能输出 JSON：
{{
  "useful": true,
  "refined_text": "string"
}}

判断与提炼规则：
1. useful 表示这段内容是否对后续 PPT 策划、结构设计、事实依据、案例、数据、结论、方案步骤、背景说明有直接帮助。
2. 如果只是冗长铺垫、重复表述、无关目录、格式噪音、版权声明、空泛客套、与用户意图明显无关的内容，应返回 useful=false。
3. 如果 useful=true，refined_text 必须提炼成适合 PPT 使用的关键信息摘要，而不是原文复读。
4. refined_text 要优先保留：主题、背景、问题、目标、结论、数据、方法、步骤、对比、案例、建议、专业术语。
5. refined_text 必须尊重用户当前需求，不能把文档主题带偏；如果与用户要求冲突，按用户要求取舍。
6. refined_text 使用中文，优先简洁但不能为了压缩长度丢失关键事实、数据、结论或步骤；通常控制在 {MAX_REFINED_CHUNK_CHARS} 字以内，必要时可以略微超出。
7. 如果 useful=false，refined_text 返回空字符串。
8. 不要输出 markdown，不要解释。

[User Requirement]
{user_requirement or "用户未提供额外文字要求。"}

[Source File]
{source.name}

[Source Chunk]
chunk {chunk_index}/{total_chunks}
{chunk}
""".strip()
        system_prompt = "你负责筛选文档中真正有助于生成 PPT 的内容，并提炼成高密度摘要。"
        raw = chat_completion_text(
            provider=runtime_cfg.text_provider,
            base_url=runtime_cfg.text_base_url,
            api_key=runtime_cfg.text_api_key,
            model=runtime_cfg.text_model,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            temperature=0.1,
        )
        parsed = self._parse_json_object(raw)
        useful = bool(parsed.get("useful"))
        refined_text = str(parsed.get("refined_text") or "").strip()
        if not useful or not refined_text:
            return ""
        return f"[Chunk {chunk_index}] {refined_text}"

    @staticmethod
    def _chunk_text(text: str, chunk_size: int) -> list[str]:
        normalized = (text or "").strip()
        if not normalized:
            return []

        chunks: list[str] = []
        current: list[str] = []
        current_len = 0
        paragraphs = [part.strip() for part in normalized.splitlines() if part.strip()]

        for paragraph in paragraphs:
            paragraph_len = len(paragraph)
            if paragraph_len >= chunk_size:
                if current:
                    chunks.append("\n".join(current).strip())
                    current = []
                    current_len = 0
                for start in range(0, paragraph_len, chunk_size):
                    piece = paragraph[start : start + chunk_size].strip()
                    if piece:
                        chunks.append(piece)
                continue

            projected = current_len + paragraph_len + (1 if current else 0)
            if current and projected > chunk_size:
                chunks.append("\n".join(current).strip())
                current = [paragraph]
                current_len = paragraph_len
            else:
                current.append(paragraph)
                current_len = projected

        if current:
            chunks.append("\n".join(current).strip())
        return [chunk for chunk in chunks if chunk]

    @staticmethod
    def _parse_json_object(raw_text: str) -> dict[str, Any]:
        text = (raw_text or "").strip()
        if text.startswith("```"):
            parts = [part for part in text.split("```") if part.strip()]
            if parts:
                text = parts[0].strip()
                if text.lower().startswith("json"):
                    text = text[4:].strip()
        start = text.find("{")
        end = text.rfind("}")
        if start < 0 or end <= start:
            raise ValueError("Model did not return valid JSON.")
        return json.loads(text[start : end + 1])
